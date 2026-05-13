# app/services/document_processor.py

import fitz  # PyMuPDF library for PDF processing 


from docx import Document # python-docx library for DOCX processing

# For OCR processing (if needed  for scanned PDFs)
import pytesseract
# from PIL import Image

from pdf2image import convert_from_path

# Regular expression for cleaning extracted text
import re

import os


# =========================
# PDF TEXT EXTRACTION
# =========================

def extract_text_from_pdf(file_path: str) -> str:

    extracted_text = ""

    pdf_document = fitz.open(file_path)

    for page in pdf_document:

        text = page.get_text()

        extracted_text += text + "\n"

    pdf_document.close()

    return extracted_text


# =========================
# DOCX TEXT EXTRACTION
# =========================

def extract_text_from_docx(file_path: str) -> str:

    document = Document(file_path)

    extracted_text = ""

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            extracted_text += text + "\n"

    return extracted_text

# =========================
# OCR FOR SCANNED PDFs
# =========================

def extract_text_with_ocr(file_path: str) -> str:

    extracted_text = ""

    # Convert PDF pages to images
    images = convert_from_path(file_path)

    # Run OCR on each page
    for image in images:

        text = pytesseract.image_to_string(image)

        extracted_text += text + "\n"

    return extracted_text

# =========================
# TEXT CLEANING
# =========================

def clean_extracted_text(text: str) -> str:

    # Remove extra spaces
    text = re.sub(r"[ ]+", " ", text)

    # Remove tabs
    text = re.sub(r"\t+", " ", text)

    # Remove excessive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Fix broken lines
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)

    # Remove page artifacts
    text = re.sub(r"Page \d+", "", text, flags=re.IGNORECASE)

    # Remove long separators
    text = re.sub(r"[-=]{3,}", "", text)

    # Strip leading/trailing whitespace
    text = text.strip()

    return text

# =========================
# PDF METADATA EXTRACTION
# =========================

def extract_pdf_metadata(file_path: str) -> dict:

    pdf_document = fitz.open(file_path)

    metadata = pdf_document.metadata

    data = {
        "pages": len(pdf_document),
        "title": metadata.get("title"),
        "author": metadata.get("author")
    }

    pdf_document.close()

    return data

# =========================
# DOCX METADATA EXTRACTION
# =========================

def extract_docx_metadata(file_path: str) -> dict:

    document = Document(file_path)

    properties = document.core_properties

    return {
        "pages": None,
        "title": properties.title,
        "author": properties.author
    }